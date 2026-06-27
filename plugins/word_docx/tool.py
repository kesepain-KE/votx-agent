"""Word DOCX 工具 — 创建 / 读取 .docx 文件，支持完整排版（字体/段落/页面/表格/图片/模板/页码/目录）。

基于 python-docx 库。table_data 和 images 接受原生对象/数组（非 JSON 字符串），
同时兼容旧版 JSON 字符串传参。

错误码: INVALID_OUTPUT_DIR / FILE_EXISTS / INVALID_COLOR / INVALID_PAGE_SIZE /
        INVALID_IMAGE_PATH / INVALID_TEMPLATE_PATH / INVALID_JSON / PDF_EXPORT_FAILED /
        RENDER_CHECK_FAILED / STRICT_MODE_ERROR / PATH_TRAVERSAL
"""
from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
from pathlib import Path

from run.tool import register_tool
from plugins._common import err, truncate, safe_path, check_sandbox
from plugins._common.artifacts import make_file_artifact, make_tool_result

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ──────── 常量 ────────

_PAGE_SIZES: dict[str, tuple[float, float]] = {
    "a4": (21.0, 29.7), "a3": (29.7, 42.0), "a5": (14.8, 21.0),
    "letter": (21.59, 27.94), "legal": (21.59, 35.56), "b5": (17.6, 25.0),
}
_ALIGN_MAP: dict[str, WD_ALIGN_PARAGRAPH] = {
    "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_LINE_SPACING_MAP: dict[str, int] = {
    "single": WD_LINE_SPACING.SINGLE,
    "1.5": WD_LINE_SPACING.ONE_POINT_FIVE,
    "double": WD_LINE_SPACING.DOUBLE,
}
_ORIENT_MAP: dict[str, WD_ORIENT] = {
    "portrait": WD_ORIENT.PORTRAIT, "landscape": WD_ORIENT.LANDSCAPE,
}
# 预设表样式
_TABLE_STYLES: set[str] = {"Table Grid", "Table Grid Light", "Light Shading", "Medium Shading 1",
                            "Medium Shading 2", "Light List", "Light Grid", "Colorful Grid",
                            "Colorful List", "Colorful Shading"}
_PAGE_NUMBER_FORMATS: dict[str, str] = {
    "1": "decimal", "第 1 页": "chineseCounting", "第 1 页 / 共 N 页": "full",
}
_RESULT_TRUNCATE = int(os.environ.get("DOCX_RESULT_TRUNCATE", "8000"))

# ──────── 样式预设 ────────

_STYLE_PRESETS: dict[str, dict] = {
    "default": {},
    "academic": {
        "font_name": "SimSun", "font_size": 12, "text_align": "justify",
        "first_line_indent": 0.74, "line_spacing": "double",
        "page_size": "a4", "margin_top": 2.54, "margin_bottom": 2.54,
        "margin_left": 3.17, "margin_right": 3.17,
        "title_font_size": 16, "title_bold": True,
        "page_number": True, "page_number_format": "1", "page_number_align": "center",
    },
    "report": {
        "font_name": "微软雅黑", "font_size": 12, "text_align": "justify",
        "first_line_indent": 0.74, "line_spacing": "1.5",
        "page_size": "a4", "margin_top": 2.54, "margin_bottom": 2.54,
        "margin_left": 3, "margin_right": 2.5,
        "title_font_size": 22, "title_bold": True, "title_spacing_after": 12,
        "page_number": True, "page_number_format": "第 1 页 / 共 N 页",
        "page_number_align": "center",
    },
    "contract": {
        "font_name": "SimSun", "font_size": 12, "text_align": "justify",
        "line_spacing": "1.5", "page_size": "a4",
        "margin_top": 2.54, "margin_bottom": 2.54,
        "margin_left": 3, "margin_right": 3,
        "title_font_size": 16, "title_bold": True,
        "page_number": True, "page_number_format": "1", "page_number_align": "center",
    },
}

# 各参数默认值（style_preset 未覆盖时使用）
_DEFAULTS: dict[str, any] = {
    "filename": "文档", "font_size": 12, "page_size": "a4", "orientation": "portrait",
    "margin_top": 2.54, "margin_bottom": 2.54, "margin_left": 2.54, "margin_right": 2.54,
    "line_spacing": "1.5", "text_align": "justify",
    "first_line_indent": 0.0, "title_font_size": 16, "title_bold": True,
    "keep_aspect_ratio": True, "auto_rename": True, "overwrite": False,
    "content_format": "plain", "table_autofit": True, "table_repeat_header": True,
    "page_number_align": "center", "page_number_format": "1",
}

# ──────── 错误码 ────────

def _ec(code: str, msg: str) -> str:
    """统一错误格式: ERROR: [CODE] message"""
    return f"ERROR: [{code}] {msg}"


def _parse_color(raw: str) -> RGBColor | None:
    """解析颜色：支持 #RRGGBB / r,g,b / 预设名。"""
    raw = (raw or "").strip().lower()
    if not raw:
        return None
    preset: dict[str, tuple[int, int, int]] = {
        "red": (255, 0, 0), "green": (0, 128, 0), "blue": (0, 0, 255),
        "black": (0, 0, 0), "white": (255, 255, 255), "gray": (128, 128, 128),
        "orange": (255, 165, 0), "purple": (128, 0, 128), "navy": (0, 0, 128),
    }
    if raw in preset:
        return RGBColor(*preset[raw])
    if raw.startswith("#") and len(raw) == 7:
        try:
            return RGBColor(int(raw[1:3], 16), int(raw[3:5], 16), int(raw[5:7], 16))
        except ValueError:
            return None
    parts = raw.split(",")
    if len(parts) == 3:
        try:
            return RGBColor(*[int(p.strip()) for p in parts])
        except ValueError:
            return None
    return None


def _parse_float(v, default: float = 0.0) -> float:
    try:
        return float(v)
    except (TypeError, ValueError):
        return default


def _parse_int(v, default: int = 0) -> int:
    try:
        return int(v)
    except (TypeError, ValueError):
        return default


# ──────── 结构化参数规范化（兼容 JSON string + 原生类型） ────────

def _normalize_table_data(raw) -> list[dict] | str | None:
    """将 table_data 规范化为 list[dict]，返回 None 表示无数据，返回 str 表示错误消息。"""
    if raw is None:
        return None
    if isinstance(raw, str):
        raw = raw.strip()
        if not raw:
            return None
        try:
            raw = json.loads(raw)
        except json.JSONDecodeError as e:
            return f"table_data 不是合法 JSON: {e}"
    if isinstance(raw, dict):
        raw = [raw]
    if not isinstance(raw, list):
        return f"table_data 应为数组或对象，收到: {type(raw).__name__}"
    # 校验每项有有效的 headers/rows
    result = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        if "headers" not in item and "rows" not in item:
            continue
        result.append(item)
    return result if result else None


def _normalize_images(raw) -> list[dict] | None:
    """将 images 规范化为 list[dict]。"""
    if raw is None:
        return None
    if isinstance(raw, str):
        raw = raw.strip()
        if not raw:
            return None
        try:
            raw = json.loads(raw)
        except json.JSONDecodeError:
            return None
    if isinstance(raw, dict):
        raw = [raw]
    if not isinstance(raw, list):
        return None
    return [i for i in raw if isinstance(i, dict) and i.get("path")]


# ──────── 轻量 Markdown → docx 渲染 ────────

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"\*(.+?)\*")
_HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
_UL_RE = re.compile(r"^[-*]\s+(.+)$")
_OL_RE = re.compile(r"^\d+[.)]\s+(.+)$")


def _render_md_paragraph(doc, text: str, font_name: str = "", font_size: int = 0,
                         font_color: str = "", text_align: str = "", content_format: str = "plain"):
    """将一段 Markdown 文本渲染为 docx 段落（仅 content_format='markdown' 时解析标记）。"""
    stripped = text.strip()
    if not stripped:
        return

    use_md = (content_format == "markdown")

    # 水平线
    if stripped in ("---", "***"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        _add_bottom_border(p)
        return

    # 标题
    if use_md:
        h_match = _HEADING_RE.match(stripped)
        if h_match:
            level = min(len(h_match.group(1)), 4)
            heading = doc.add_heading(h_match.group(2), level=level)
            if font_color:
                c = _parse_color(font_color)
                if c:
                    for run in heading.runs:
                        run.font.color.rgb = c
            return

        # 无序列表
        ul_match = _UL_RE.match(stripped)
        if ul_match:
            p = doc.add_paragraph(style='List Bullet')
            _render_inline_runs(p, ul_match.group(1), font_name, font_size, font_color)
            return

        # 有序列表
        ol_match = _OL_RE.match(stripped)
        if ol_match:
            p = doc.add_paragraph(style='List Number')
            _render_inline_runs(p, ol_match.group(1), font_name, font_size, font_color)
            return

    # 普通段落
    p = doc.add_paragraph()
    if text_align in _ALIGN_MAP:
        p.alignment = _ALIGN_MAP[text_align]
    p.paragraph_format.line_spacing = 1.15

    if use_md:
        _render_inline_runs(p, stripped, font_name, font_size, font_color)
    else:
        run = p.add_run(stripped)
        if font_name:
            run.font.name = font_name
            _set_east_asian_font(run, font_name)
        if font_size > 0:
            run.font.size = Pt(font_size)
        if font_color:
            c = _parse_color(font_color)
            if c:
                run.font.color.rgb = c


def _render_inline_runs(p, text: str, font_name: str = "", font_size: int = 0, font_color: str = ""):
    """在段落中逐个添加 run，处理 **加粗** 和 *斜体*。"""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            run = p.add_run(inner)
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            inner = part[1:-1]
            run = p.add_run(inner)
            run.italic = True
        else:
            run = p.add_run(part)

        if font_name:
            run.font.name = font_name
            _set_east_asian_font(run, font_name)
        if font_size > 0:
            run.font.size = Pt(font_size)
        if font_color:
            c = _parse_color(font_color)
            if c:
                run.font.color.rgb = c


def _set_east_asian_font(run, font_name: str):
    """设置东亚字体回退（中/日/韩）。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _add_bottom_border(paragraph):
    """给段落添加底部边框（水平线）。"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ──────── create_docx ────────

def create_docx(
    output_dir: str = "",
    output_path: str = "",
    filename: str = "",
    title: str = "",
    content: str = "",
    overwrite: bool = False,
    auto_rename: bool = True,
    # ── 内容格式 ──
    content_format: str = "plain",
    text_align: str = "",
    # ── 字体 (正文) ──
    font_name: str = "",
    font_size: int = 0,
    font_color: str = "",
    # ── 标题样式 ──
    title_font_size: int = 0,
    title_bold: bool = True,
    title_spacing_after: float = 0,
    # ── 页面 ──
    page_size: str = "",
    orientation: str = "",
    margin_top: float = 0,
    margin_bottom: float = 0,
    margin_left: float = 0,
    margin_right: float = 0,
    # ── 段落 ──
    line_spacing: str = "",
    paragraph_spacing: float = 0,
    first_line_indent: float = 0,
    # ── 页眉页脚 / 页码 ──
    header_text: str = "",
    footer_text: str = "",
    page_number: bool = False,
    page_number_format: str = "",
    page_number_align: str = "",
    # ── 表格 ──
    table_data=None,  # list[dict] | dict | JSON string
    table_style: str = "",
    table_autofit: bool = True,
    table_repeat_header: bool = True,
    # ── 图片 ──
    images=None,  # list[dict] | dict | JSON string
    keep_aspect_ratio: bool = True,
    # ── 模板 ──
    template_path: str = "",
    # ── 工程化 ──
    export_pdf: bool = False,
    render_check: bool = False,
    strict_mode: bool = False,
    # ── 预设 / 元信息 ──
    style_preset: str = "",
    toc: bool = False,
    metadata: str = "",
    language: str = "",
) -> str:
    """创建 Word 文档（.docx），约 43 个实际入参，覆盖 20 类排版/工程化能力。

    参数优先级: style_preset 预设 < 显式传入参数 < strict_mode 校验。
    content 当 content_format='markdown' 时支持：# 标题、**加粗**、*斜体*、- 列表、--- 水平线。
    table_data/images 接受原生 list/dict 对象（也兼容旧版 JSON 字符串）。

    错误码: INVALID_OUTPUT_DIR / FILE_EXISTS / INVALID_COLOR / INVALID_PAGE_SIZE /
            INVALID_IMAGE_PATH / INVALID_TEMPLATE_PATH / INVALID_TABLE_DATA /
            INVALID_JSON / RENDER_CHECK_FAILED / STRICT_MODE_ERROR / PATH_TRAVERSAL /
            SAVE_FAILED / DEPENDENCY
    """
    if not HAS_DOCX:
        return _ec("DEPENDENCY", "python-docx 未安装，请执行: pip install python-docx")

    # ── style_preset 预设（先于显式参数应用） ──
    sp = (style_preset or "").strip().lower()
    preset = _STYLE_PRESETS.get(sp, {})
    if sp and not preset and strict_mode:
        return _ec("STRICT_MODE_ERROR", f"style_preset 不支持: {sp}，可选: {', '.join(_STYLE_PRESETS)}")
    # 用预设覆盖未显式传入的参数
    _apply_preset = lambda key, current: preset.get(key) if not current and preset.get(key) is not None else current
    overwrite = _apply_preset("overwrite", overwrite)
    auto_rename = _apply_preset("auto_rename", auto_rename)
    font_name = _apply_preset("font_name", font_name)
    font_size = _apply_preset("font_size", font_size) if font_size else (_DEFAULTS["font_size"] if not preset else preset.get("font_size", font_size))
    text_align = _apply_preset("text_align", text_align)
    first_line_indent = _apply_preset("first_line_indent", first_line_indent) if first_line_indent else (preset.get("first_line_indent", first_line_indent))
    line_spacing = _apply_preset("line_spacing", line_spacing)
    page_size = _apply_preset("page_size", page_size)
    margin_top = _apply_preset("margin_top", margin_top) if margin_top else (preset.get("margin_top", margin_top))
    margin_bottom = _apply_preset("margin_bottom", margin_bottom) if margin_bottom else (preset.get("margin_bottom", margin_bottom))
    margin_left = _apply_preset("margin_left", margin_left) if margin_left else (preset.get("margin_left", margin_left))
    margin_right = _apply_preset("margin_right", margin_right) if margin_right else (preset.get("margin_right", margin_right))
    title_font_size = _apply_preset("title_font_size", title_font_size) if title_font_size else (preset.get("title_font_size", title_font_size))
    title_bold = _apply_preset("title_bold", title_bold) if isinstance(title_bold, bool) else (preset.get("title_bold", title_bold))
    title_spacing_after = _apply_preset("title_spacing_after", title_spacing_after) if title_spacing_after else (preset.get("title_spacing_after", title_spacing_after))
    page_number = _apply_preset("page_number", page_number)
    page_number_format = _apply_preset("page_number_format", page_number_format)
    page_number_align = _apply_preset("page_number_align", page_number_align)

    # ── 输出目录：output_dir 优先，output_path 兜底（向后兼容） ──
    od = (output_dir or output_path or "").strip()
    if not od:
        return _ec("INVALID_OUTPUT_DIR", "output_dir 不能为空")
    try:
        out_p = safe_path(od)
        if out_p is None:
            return err(f"路径无效: {od}")
        if check_sandbox(out_p) is None:
            return _ec("PATH_TRAVERSAL", f"路径越权（仅允许项目目录和用户目录）: {od}")
        out_p.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        return _ec("INVALID_OUTPUT_DIR", f"输出目录创建失败: {e}")

    # ── 文件名冲突处理 ──
    name = (filename or "文档").strip()
    if not name.endswith(".docx"):
        name += ".docx"
    filepath = out_p / name

    if filepath.exists():
        if overwrite:
            pass
        elif auto_rename:
            stem = filepath.stem
            suffix = filepath.suffix
            for i in range(1, 1000):
                candidate = out_p / f"{stem}_{i}{suffix}"
                if not candidate.exists():
                    filepath = candidate
                    break
            else:
                return _ec("FILE_EXISTS", f"无法生成不冲突的文件名: {filepath}")
        else:
            return _ec("FILE_EXISTS", f"文件已存在: {filepath}（设置 overwrite=true 或 auto_rename=true）")

    # ── 模板 ──
    tp = (template_path or "").strip()
    if tp:
        tpp = safe_path(tp)
        if tpp is None or check_sandbox(tpp) is None:
            return _ec("INVALID_TEMPLATE_PATH", f"模板路径越权: {tp}")
        if not tpp.exists():
            return _ec("INVALID_TEMPLATE_PATH", f"模板文件不存在: {tp}")
        if not tpp.suffix.lower() == ".docx":
            return _ec("INVALID_TEMPLATE_PATH", f"模板不是 .docx 文件: {tp}")
        try:
            doc = Document(str(tpp))
        except Exception as e:
            return _ec("INVALID_TEMPLATE_PATH", f"无法打开模板文件: {e}")
    else:
        doc = Document()

    # ── metadata / language ──
    lang = (language or "").strip()
    if lang:
        try:
            doc.element.set(qn('xml:lang'), lang)
        except Exception:
            if strict_mode:
                return _ec("STRICT_MODE_ERROR", f"language 设置失败: {lang}")
    meta = (metadata or "").strip()
    if meta:
        try:
            md = json.loads(meta) if isinstance(meta, str) else meta
            if isinstance(md, dict):
                cp = doc.core_properties
                if md.get("author"): cp.author = str(md["author"])
                if md.get("subject"): cp.subject = str(md["subject"])
                if md.get("keywords"): cp.keywords = str(md["keywords"])
                if md.get("title"): cp.title = str(md["title"])
        except json.JSONDecodeError:
            if strict_mode:
                return _ec("INVALID_JSON", f"metadata 不是合法 JSON: {meta}")
        except Exception:
            if strict_mode:
                return _ec("STRICT_MODE_ERROR", "metadata 写入失败")

    # ── 页面设置 ──
    section = doc.sections[0]
    ps = (page_size or "").strip().lower()
    if ps in _PAGE_SIZES:
        w, h = _PAGE_SIZES[ps]
        section.page_width = Cm(w)
        section.page_height = Cm(h)
    elif ps and strict_mode:
        return _ec("INVALID_PAGE_SIZE", f"page_size 不支持: {ps}，可选: {', '.join(_PAGE_SIZES)}")
    ori = (orientation or "").strip().lower()
    if ori in _ORIENT_MAP:
        section.orientation = _ORIENT_MAP[ori]
        if ori == "landscape":
            section.page_width, section.page_height = section.page_height, section.page_width
    if margin_top > 0:
        section.top_margin = Cm(margin_top)
    if margin_bottom > 0:
        section.bottom_margin = Cm(margin_bottom)
    if margin_left > 0:
        section.left_margin = Cm(margin_left)
    if margin_right > 0:
        section.right_margin = Cm(margin_right)

    # ── 默认字体 ──
    style = doc.styles['Normal']
    if font_name:
        style.font.name = font_name
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size > 0:
        style.font.size = Pt(font_size)
    if font_color:
        c = _parse_color(font_color)
        if c:
            style.font.color.rgb = c
        elif strict_mode:
            return _ec("INVALID_COLOR", f"font_color 格式无效: {font_color}（支持: #RRGGBB / r,g,b / red 等预设名）")
    if paragraph_spacing > 0:
        style.paragraph_format.space_after = Pt(paragraph_spacing)
    if line_spacing:
        ls = line_spacing.strip().lower()
        if ls in _LINE_SPACING_MAP:
            style.paragraph_format.line_spacing = _LINE_SPACING_MAP[ls]
        else:
            val = _parse_float(line_spacing, 0)
            if val > 0:
                style.paragraph_format.line_spacing = val
    if first_line_indent > 0:
        style.paragraph_format.first_line_indent = Cm(first_line_indent)

    # ── 页眉 ──
    header_text = (header_text or "").strip()
    if header_text:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = header_text
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 页脚（页脚文字 + 页码） ──
    has_footer = bool((footer_text or "").strip()) or page_number
    if has_footer:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = _ALIGN_MAP.get(page_number_align, WD_ALIGN_PARAGRAPH.CENTER)

        ft = (footer_text or "").strip()
        if ft and page_number:
            # 页脚文字 + 页码
            fmt = (page_number_format or "1").strip()
            _add_page_number(fp, fmt, prefix=ft + "  ")
        elif page_number:
            fmt = (page_number_format or "1").strip()
            _add_page_number(fp, fmt)
        else:
            fp.text = ft

    # ── 文档标题 ──
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title_font_size > 0:
            for run in heading.runs:
                run.font.size = Pt(title_font_size)
        if not title_bold:
            for run in heading.runs:
                run.bold = False
        if title_spacing_after > 0:
            heading.paragraph_format.space_after = Pt(title_spacing_after)

    # ── 目录（TOC） ──
    if toc:
        toc_para = doc.add_paragraph()
        toc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = toc_para.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        run2 = toc_para.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run2._element.append(instrText)
        run3 = toc_para.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run3._element.append(fldChar2)
        run4 = toc_para.add_run('[目录 — 请在 Word 中右键点击「更新域」以生成]')
        run4.font.color.rgb = RGBColor(128, 128, 128)
        run4.font.size = Pt(10)
        run5 = toc_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run5._element.append(fldChar3)
        # 分页
        doc.add_page_break()

    # ── 正文渲染 ──
    cf = (content_format or "plain").strip().lower()
    if cf not in ("plain", "markdown"):
        cf = "plain"
    if content:
        for line in content.split("\n"):
            _render_md_paragraph(
                doc, line,
                font_name=font_name, font_size=font_size, font_color=font_color,
                text_align=text_align, content_format=cf,
            )

    # ── 表格 ──
    tables = _normalize_table_data(table_data)
    if isinstance(tables, str):
        if strict_mode:
            return _ec("INVALID_TABLE_DATA", tables)
        tables = None
    if tables is not None and not isinstance(tables, list):
        if strict_mode:
            return _ec("INVALID_TABLE_DATA", f"table_data 格式错误：应为数组或对象。收到: {type(table_data).__name__}")
        tables = None
    if isinstance(tables, list) and tables:
        ts = (table_style or "").strip()
        if ts and ts in _TABLE_STYLES:
            pass  # 使用用户指定样式
        elif ts:
            pass  # 尝试自定义样式名
        else:
            ts = "Table Grid"  # 默认
        _insert_tables(doc, tables, font_name, font_size, font_color,
                       table_style=ts, autofit=table_autofit,
                       repeat_header=table_repeat_header)

    # ── 图片 ──
    imgs = _normalize_images(images)
    if isinstance(imgs, list) and imgs:
        _insert_images(doc, imgs, keep_aspect_ratio=keep_aspect_ratio)

    # ── 保存 ──
    try:
        doc.save(str(filepath))
    except Exception as e:
        return _ec("SAVE_FAILED", f"保存文档失败: {e}")

    artifacts = [make_file_artifact(filepath)]
    warnings: list[str] = []
    extra: dict[str, object] = {}

    if render_check:
        issues = _render_check(filepath)
        if issues:
            warnings.extend(issues)
            extra["render_check"] = {"passed": False, "issues": issues}
        else:
            extra["render_check"] = {"passed": True}

    if export_pdf:
        pdf_ok = False
        pdf_path = filepath.with_suffix(".pdf")
        office_cmd = shutil.which("libreoffice") or shutil.which("soffice")
        try:
            if office_cmd:
                r = subprocess.run(
                    [office_cmd, "--headless", "--convert-to", "pdf", str(filepath)],
                    cwd=str(filepath.parent), capture_output=True, timeout=60,
                )
                pdf_ok = r.returncode == 0 and pdf_path.exists()
        except Exception:
            pass
        if pdf_ok:
            artifacts.append(make_file_artifact(pdf_path))
            extra["export_pdf"] = {"success": True, "path": str(pdf_path)}
        else:
            warnings.append(
                "转 PDF 失败：需要 LibreOffice/soffice。"
                f" 可手动执行: soffice --headless --convert-to pdf \"{filepath}\""
            )
            extra["export_pdf"] = {"success": False}

    if warnings:
        extra["warnings"] = warnings

    return make_tool_result(True, "文档生成完成", artifacts, **extra)

    size = filepath.stat().st_size
    rel_path = str(filepath)
    lines = [
        f"[SUCCESS]",
        f"  docx: {rel_path}",
        f"  size: {size} bytes ({size / 1024:.1f} KB)",
    ]

    # ── 警告 ──
    warnings: list[str] = []

    # render_check
    if render_check:
        issues = _render_check(filepath)
        if issues:
            warnings.extend(issues)
            lines.append(f"  render: FAILED ({len(issues)} issues)")
            for issue in issues:
                lines.append(f"    - {issue}")
        else:
            lines.append("  render: passed")

    # export_pdf
    if export_pdf:
        pdf_ok = False
        office_cmd = shutil.which("libreoffice") or shutil.which("soffice")
        try:
            if office_cmd:
                r = subprocess.run(
                    [office_cmd, "--headless", "--convert-to", "pdf", str(filepath)],
                    cwd=str(filepath.parent), capture_output=True, timeout=60,
                )
                pdf_ok = r.returncode == 0
        except Exception:
            pass
        if pdf_ok:
            pdf_path = str(filepath.with_suffix(".pdf"))
            lines.append(f"  pdf: {pdf_path}")
        else:
            lines.append(f"  pdf: FAILED（转 PDF 需 LibreOffice/soffice。可手动: soffice --headless --convert-to pdf \"{rel_path}\"）")

    if warnings:
        lines.append(f"  warnings: {len(warnings)}")
    else:
        lines.append("  warnings: none")

    return "\n".join(lines)


def _add_page_number(paragraph, fmt: str, prefix: str = ""):
    """在段落中插入页码域。fmt: '1' | '第 1 页' | '第 1 页 / 共 N 页'"""
    if prefix:
        run = paragraph.add_run(prefix)
        run.font.size = Pt(9)

    if fmt == "第 1 页 / 共 N 页":
        run = paragraph.add_run("第 ")
        run.font.size = Pt(9)
        _add_field(paragraph, "PAGE")
        run = paragraph.add_run(" 页 / 共 ")
        run.font.size = Pt(9)
        _add_field(paragraph, "NUMPAGES")
        run = paragraph.add_run(" 页")
        run.font.size = Pt(9)
    elif fmt == "第 1 页":
        run = paragraph.add_run("第 ")
        run.font.size = Pt(9)
        _add_field(paragraph, "PAGE")
        run = paragraph.add_run(" 页")
        run.font.size = Pt(9)
    else:
        _add_field(paragraph, "PAGE")


def _add_field(paragraph, field_type: str):
    """插入 Word 域（PAGE / NUMPAGES 等）。"""
    run = paragraph.add_run()
    run.font.size = Pt(9)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_type
    run2 = paragraph.add_run()
    run2._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3 = paragraph.add_run()
    run3._element.append(fldChar2)


def _render_check(filepath: Path) -> list[str]:
    """基本渲染检查：重新打开文档，检测常见问题。"""
    issues: list[str] = []
    try:
        doc = Document(str(filepath))
    except Exception as e:
        return [f"无法重新打开文档: {e}"]

    # 检查正文是否为空
    total_text = "".join(p.text for p in doc.paragraphs)
    if len(total_text.strip()) < 10 and not doc.tables:
        issues.append("文档正文过短（可能为空）")

    # 检查表格溢出（列数过多）
    for ti, table in enumerate(doc.tables):
        ncols = len(table.columns)
        if ncols > 8:
            issues.append(f"表格 {ti + 1} 有 {ncols} 列，可能横向溢出")

    # 检查图片数量
    img_count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            for _ in run._element.findall(qn('w:drawing')):
                img_count += 1
            for _ in run._element.findall(qn('wp:inline')):
                img_count += 1
    if img_count == 0:
        # 不做强制警告——不是所有文档都需要图片
        pass

    return issues


# ──────── 表格 ────────

def _insert_tables(doc, tables: list[dict], font_name: str, font_size: int, font_color: str,
                   table_style: str = "Table Grid", autofit: bool = True,
                   repeat_header: bool = True):
    """插入表格。每项: {"caption", "headers": [...], "rows": [[...], ...]}"""
    for ti, tdef in enumerate(tables):
        if not isinstance(tdef, dict):
            continue
        headers = tdef.get("headers", [])
        rows = tdef.get("rows", [])
        caption = tdef.get("caption", "")

        if not headers and not rows:
            continue

        ncols = len(headers) if headers else (len(rows[0]) if rows and rows[0] else 1)
        table = doc.add_table(rows=1 + len(rows), cols=ncols)

        # 表样式
        try:
            table.style = table_style
        except Exception:
            pass  # 样式不存在则使用默认

        if autofit:
            table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        if headers:
            hdr_row = table.rows[0]
            if repeat_header:
                # 跨页重复表头
                trPr = hdr_row._tr.get_or_add_trPr()
                tblHeader = OxmlElement('w:tblHeader')
                tblHeader.set(qn('w:val'), 'true')
                trPr.append(tblHeader)

            for ci, h in enumerate(headers):
                if ci < ncols:
                    cell = hdr_row.cells[ci]
                    cell.text = str(h)
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True
                            if font_name:
                                run.font.name = font_name
                                _set_east_asian_font(run, font_name)
                            if font_size > 0:
                                run.font.size = Pt(font_size)
                    _set_cell_shading(cell, "D9D9D9")

        # 数据行
        for ri, row in enumerate(rows):
            if not isinstance(row, list):
                continue
            for ci, val in enumerate(row):
                if ci < ncols:
                    cell = table.rows[1 + ri].cells[ci]
                    cell.text = str(val) if val is not None else ""
                    if font_name:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.name = font_name
                                _set_east_asian_font(run, font_name)

        # 题注
        if caption:
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(f"表 {ti + 1}: {caption}" if len(tables) > 1 else caption)
            cap_run.bold = True
            cap_run.font.size = Pt(10)
            cap_para.paragraph_format.space_after = Pt(6)


def _set_cell_shading(cell, color: str):
    """单元格背景色。"""
    tcPr = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


# ──────── 图片 ────────

def _insert_images(doc, imgs: list[dict], keep_aspect_ratio: bool = True):
    """插入图片。每项: {"path", "width", "height", "alignment", "caption"}"""
    for img_def in imgs:
        if not isinstance(img_def, dict):
            continue
        img_path = img_def.get("path", "")
        if not img_path:
            continue
        pp = safe_path(img_path)
        if pp is None or check_sandbox(pp) is None:
            continue
        if not pp.exists():
            continue

        # 图片段落
        p = doc.add_paragraph()
        align = (img_def.get("alignment") or "center").strip().lower()
        if align in _ALIGN_MAP:
            p.alignment = _ALIGN_MAP[align]

        width = _parse_int(img_def.get("width"), 0)
        height = _parse_int(img_def.get("height"), 0)

        run = p.add_run()
        try:
            if width > 0 and height > 0:
                run.add_picture(str(pp), width=Pt(width), height=Pt(height))
            elif width > 0:
                run.add_picture(str(pp), width=Pt(width))
            elif height > 0:
                run.add_picture(str(pp), height=Pt(height))
            else:
                run.add_picture(str(pp), width=Inches(5))
        except Exception:
            run.text = f"[图片: {pp.name}]"

        # 题注
        caption = (img_def.get("caption") or "").strip()
        if caption:
            cap_para = doc.add_paragraph()
            cap_para.alignment = p.alignment
            cap_run = cap_para.add_run(caption)
            cap_run.font.size = Pt(10)
            cap_run.italic = True
            cap_para.paragraph_format.space_after = Pt(10)


# ──────── read_docx ────────

def read_docx(path: str) -> str:
    """读取 Word 文档（含段落样式/格式/表格/图片/页面信息）。"""
    if not HAS_DOCX:
        return err("python-docx 未安装，请执行: pip install python-docx")

    try:
        pp = safe_path(path)
        if pp is None:
            return err(f"路径无效: {path}")
        if check_sandbox(pp) is None:
            return err(f"路径越权（仅允许项目目录和用户目录）: {path}")
        if not pp.exists():
            return err(f"文件不存在: {pp}")
        if not pp.suffix.lower() == ".docx":
            return err(f"不是 .docx 文件: {pp}")

        doc = Document(str(pp))
        lines = [f"📄 {pp.name}"]

        # 页面
        sections = doc.sections
        if sections:
            s = sections[0]
            pw_inch = s.page_width / 914400
            ph_inch = s.page_height / 914400
            lines.append(f"📐 页面: {pw_inch:.1f}x{ph_inch:.1f} inch | "
                         f"边距 L{s.left_margin / 914400:.1f} R{s.right_margin / 914400:.1f} inch")
            lines.append(f"   方向: {'横向' if s.orientation == WD_ORIENT.LANDSCAPE else '纵向'}")

        # 统计
        img_count = sum(len(run._element.findall(qn('w:drawing'))) +
                        len(run._element.findall(qn('wp:inline')))
                        for p in doc.paragraphs for run in p.runs)
        lines.append(f"📊 段落: {len(doc.paragraphs)} | 表格: {len(doc.tables)}" +
                     (f" | 图片: {img_count}" if img_count else ""))
        lines.append("")

        # 段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text and not para.runs:
                continue
            style_name = para.style.name if para.style else "Normal"
            info: list[str] = []
            if para.runs:
                r = para.runs[0]
                if r.bold: info.append("B")
                if r.italic: info.append("I")
                if r.underline: info.append("U")
                if r.font.size: info.append(f"{r.font.size / 12700:.0f}pt")
                if r.font.color and r.font.color.rgb:
                    info.append(f"#{r.font.color.rgb}")
            prefix = f"[{style_name}]"
            if info:
                prefix += f" ({','.join(info)})"
            lines.append(f"{prefix} {text}")

        # 表格
        for ti, table in enumerate(doc.tables):
            lines.append(f"\n── 表格 {ti + 1} ({len(table.rows)}行 x {len(table.columns)}列) ──")
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                lines.append(f"  [{ri + 1}] {' | '.join(cells)}")

        return truncate("\n".join(lines), _RESULT_TRUNCATE)

    except Exception as e:
        return err(f"读取文档失败: {e}")


# ──────── 注册 ────────

SCHEMAS = [
    {
        "type": "function",
        "function": {
            "name": "create_docx",
            "description": (
                "创建 Word 文档（.docx），支持完整排版：标题/正文(Markdown)/字体/字号/颜色/"
                "页面大小/边距/方向/行距/缩进/对齐/页眉页脚/页码/表格/图片/模板/目录/预设风格。"
                "~35 参数覆盖 20 类排版项。style_preset 可快速套用 academic/report/contract 预设。"
                "参数优先级: style_preset 预设 < 显式传入参数。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    # ── 基础 ──
                    "output_dir": {
                        "type": "string",
                        "description": "输出目录路径（必填）。优先于此参数，output_path 已废弃但仍兼容",
                    },
                    "output_path": {
                        "type": "string",
                        "description": "[已废弃] 输出目录路径，请改用 output_dir",
                    },
                    "filename": {
                        "type": "string",
                        "description": "文件名不含扩展名，默认「文档」",
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题（自动居中加粗）",
                    },
                    "content": {
                        "type": "string",
                        "description": (
                            "正文，多行用 \\n 分隔。当 content_format='markdown' 时支持：\n"
                            "#/##/###/#### 标题、**加粗**、*斜体*、- 或 * 无序列表、1. 有序列表、--- 水平线"
                        ),
                    },
                    "overwrite": {
                        "type": "boolean",
                        "description": "文件已存在时是否覆盖，默认 false。与 auto_rename 互斥（auto_rename 优先）",
                    },
                    "auto_rename": {
                        "type": "boolean",
                        "description": "文件重名时自动追加序号如 文档_1.docx。默认 true",
                    },
                    # ── 内容格式 ──
                    "content_format": {
                        "type": "string",
                        "enum": ["plain", "markdown"],
                        "description": "正文格式：plain 纯文本（默认），markdown 支持轻量标记",
                    },
                    "text_align": {
                        "type": "string",
                        "enum": ["left", "center", "right", "justify"],
                        "description": "正文段落对齐方式。中文正式文档建议 justify（两端对齐）",
                    },
                    # ── 字体 ──
                    "font_name": {
                        "type": "string",
                        "description": "正文字体，如 微软雅黑、SimSun、Arial",
                    },
                    "font_size": {
                        "type": "integer",
                        "description": "正文字号（磅/pt）：12=小四、10.5=五号、16=三号",
                    },
                    "font_color": {
                        "type": "string",
                        "description": "正文字体颜色：#RRGGBB、r,g,b、或预设 red/green/blue/black/gray",
                    },
                    # ── 标题样式 ──
                    "title_font_size": {
                        "type": "integer",
                        "description": "标题字号（磅/pt），独立于正文。如 22=二号、18=小二号。不传则用 Word 默认",
                    },
                    "title_bold": {
                        "type": "boolean",
                        "description": "标题是否加粗，默认 true",
                    },
                    "title_spacing_after": {
                        "type": "number",
                        "description": "标题段后间距（磅/pt），如 12=空一行",
                    },
                    # ── 页面 ──
                    "page_size": {
                        "type": "string",
                        "enum": ["a4", "a3", "a5", "letter", "legal", "b5"],
                        "description": "纸张大小，默认 A4",
                    },
                    "orientation": {
                        "type": "string",
                        "enum": ["portrait", "landscape"],
                        "description": "方向：portrait 纵向（默认）、landscape 横向",
                    },
                    "margin_top": {"type": "number", "description": "上边距（厘米）"},
                    "margin_bottom": {"type": "number", "description": "下边距（厘米）"},
                    "margin_left": {"type": "number", "description": "左边距（厘米）"},
                    "margin_right": {"type": "number", "description": "右边距（厘米）"},
                    # ── 段落 ──
                    "line_spacing": {
                        "type": "string",
                        "description": "行间距：single 单倍、1.5 倍、double 双倍，或数字如 1.15",
                    },
                    "paragraph_spacing": {
                        "type": "number",
                        "description": "段后间距（磅/pt）",
                    },
                    "first_line_indent": {
                        "type": "number",
                        "description": "首行缩进（厘米），中文常用 0.74（两个字符）",
                    },
                    # ── 页眉页脚 / 页码 ──
                    "header_text": {"type": "string", "description": "页眉文字（居中）"},
                    "footer_text": {"type": "string", "description": "页脚文字（居中）。与 page_number 可共存"},
                    "page_number": {"type": "boolean", "description": "是否显示页码，默认 false"},
                    "page_number_format": {
                        "type": "string",
                        "enum": ["1", "第 1 页", "第 1 页 / 共 N 页"],
                        "description": "页码格式，默认 '1'",
                    },
                    "page_number_align": {
                        "type": "string",
                        "enum": ["left", "center", "right"],
                        "description": "页码对齐，默认 center",
                    },
                    # ── 表格 ──
                    "table_data": {
                        "type": "array",
                        "description": (
                            "表格数据（原生数组，非 JSON 字符串）。每项：\n"
                            '{"caption":"表名","headers":["列A","列B"],"rows":[["1","2"]]}\n'
                            "headers 可选，rows 是二维数组"
                        ),
                        "items": {
                            "type": "object",
                            "properties": {
                                "caption": {"type": "string", "description": "表格题注"},
                                "headers": {
                                    "type": "array", "items": {"type": "string"},
                                    "description": "表头列名",
                                },
                                "rows": {
                                    "type": "array", "items": {"type": "array", "items": {"type": "string"}},
                                    "description": "数据行，每行是字符串数组",
                                },
                            },
                        },
                    },
                    "table_style": {
                        "type": "string",
                        "description": (
                            "表格样式名。常用：Table Grid（默认）/ Light Shading / "
                            "Medium Shading 1 / Colorful Grid / Light List 等"
                        ),
                    },
                    "table_autofit": {
                        "type": "boolean",
                        "description": "是否自动适配列宽，默认 true",
                    },
                    "table_repeat_header": {
                        "type": "boolean",
                        "description": "跨页时是否重复表头，默认 true（长表格推荐开启）",
                    },
                    # ── 图片 ──
                    "images": {
                        "type": "array",
                        "description": (
                            "图片列表（原生数组，非 JSON 字符串）。每项：\n"
                            '{"path":"/path/img.png","width":300,"height":200,"alignment":"center","caption":"图1"}'
                        ),
                        "items": {
                            "type": "object",
                            "properties": {
                                "path": {"type": "string", "description": "图片文件路径（必填）"},
                                "width": {"type": "integer", "description": "显示宽度（磅/pt）"},
                                "height": {"type": "integer", "description": "显示高度（磅/pt）"},
                                "alignment": {
                                    "type": "string", "enum": ["left", "center", "right"],
                                    "description": "对齐，默认 center",
                                },
                                "caption": {"type": "string", "description": "图片题注"},
                            },
                        },
                    },
                    "keep_aspect_ratio": {
                        "type": "boolean",
                        "description": "插入图片时是否保持宽高比。默认 true。设 false 可能导致图片压扁",
                    },
                    # ── 模板 ──
                    "template_path": {
                        "type": "string",
                        "description": "已有 .docx 模板路径。使用模板时，页面/字体参数会覆盖模板默认值",
                    },
                    # ── 工程化 ──
                    "export_pdf": {
                        "type": "boolean",
                        "description": "是否同时导出 PDF（需 LibreOffice）。失败不影响 DOCX 生成",
                    },
                    "render_check": {
                        "type": "boolean",
                        "description": "生成后渲染检查（空文档/表格溢出/图片缺失），默认 false",
                    },
                    "strict_mode": {
                        "type": "boolean",
                        "description": "参数非法时直接报错，不自动兜底。默认 false",
                    },
                    "style_preset": {
                        "type": "string",
                        "enum": ["default", "academic", "report", "contract"],
                        "description": (
                            "快速套用文档风格预设：\n"
                            "default=无预设 / academic=学术论文(宋体12pt/双倍行距/页码) / "
                            "report=中文报告(微软雅黑/1.5倍行距/首行缩进) / "
                            "contract=合同(宋体12pt/1.5倍行距)。\n"
                            "预设参数可被显式传入参数覆盖"
                        ),
                    },
                    "toc": {
                        "type": "boolean",
                        "description": (
                            "是否根据标题生成目录（TOC 域）。"
                            "需要配合 content_format='markdown' 使用 # 标题。"
                            "在 Word 中右键点击「更新域」即可生成。默认 false"
                        ),
                    },
                    "metadata": {
                        "type": "string",
                        "description": (
                            "文档元信息 JSON 字符串: "
                            '{"author":"作者","subject":"主题","keywords":"关键词","title":"文档属性标题"}'
                        ),
                    },
                    "language": {
                        "type": "string",
                        "description": "文档语言标记，如 zh-CN / en-US。影响拼写检查和断字",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": (
                "读取 Word 文档（.docx）内容，返回段落（含样式名和加粗/斜体/字号/颜色格式）、"
                "表格、图片数量、页面边距和方向信息。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": ".docx 文件路径（必填）",
                    },
                },
                "required": ["path"],
            },
        },
    },
]

HANDLERS = {
    "create_docx": create_docx,
    "read_docx": read_docx,
}


def register():
    """注册 word_docx 工具。"""
    for s in SCHEMAS:
        name = s["function"]["name"]
        register_tool(s, HANDLERS[name])
