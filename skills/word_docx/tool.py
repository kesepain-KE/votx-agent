"""Word DOCX 工具 — 创建 / 读取 .docx 文件"""
import os
from pathlib import Path
from run.tool import register_tool
from skills._common import err, truncate, safe_path, check_sandbox

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def create_docx(
    output_path: str,
    title: str = "",
    content: str = "",
    filename: str = ""
) -> str:
    """创建 Word 文档
    
    Args:
        output_path: 输出目录路径
        title: 文档标题
        content: 正文内容（支持多行）
        filename: 文件名（不含扩展名，默认"文档"）
    """
    if not HAS_DOCX:
        return err("python-docx 未安装，请执行: pip install python-docx")
    
    try:
        p = safe_path(output_path)
        if p is None:
            return err(f"路径无效: {output_path}")
        if check_sandbox(p) is None:
            return err(f"路径越权（仅允许项目目录和用户目录）: {output_path}")

        p.mkdir(parents=True, exist_ok=True)
        
        name = (filename or "文档").strip()
        if not name.endswith(".docx"):
            name += ".docx"
        
        filepath = p / name
        
        doc = Document()
        
        # 标题
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 正文
        if content:
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        
        doc.save(str(filepath))
        return f"OK: 已创建文档 {filepath} ({filepath.stat().st_size} bytes)"
    
    except Exception as e:
        return err(f"创建文档失败: {e}")


def read_docx(path: str) -> str:
    """读取 Word 文档内容
    
    Args:
        path: .docx 文件路径
    """
    if not HAS_DOCX:
        return err("python-docx 未安装，请执行: pip install python-docx")
    
    try:
        p = safe_path(path)
        if p is None:
            return err(f"路径无效: {path}")
        if check_sandbox(p) is None:
            return err(f"路径越权（仅允许项目目录和用户目录）: {path}")
        if not p.exists():
            return err(f"文件不存在: {p}")
        if not p.suffix.lower() == ".docx":
            return err(f"不是 .docx 文件: {p}")
        
        doc = Document(str(p))
        lines = [f"📄 {p.name}"]
        lines.append(f"段落数: {len(doc.paragraphs)}")
        lines.append(f"表格数: {len(doc.tables)}")
        lines.append("")
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else "Normal"
                lines.append(f"[{style}] {text}")
        
        for ti, table in enumerate(doc.tables):
            lines.append(f"\n--- 表格 {ti+1} ({len(table.rows)}行 x {len(table.columns)}列) ---")
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(f"  行{ri+1}: {' | '.join(cells)}")
        
        return truncate("\n".join(lines))
    
    except Exception as e:
        return err(f"读取文档失败: {e}")


# —— Schema 定义 ——

SCHEMAS = [
    {
        "type": "function",
        "function": {
            "name": "create_docx",
            "description": "创建 Word 文档（.docx），支持标题和正文内容",
            "parameters": {
                "type": "object",
                "properties": {
                    "output_path": {
                        "type": "string",
                        "description": "输出目录路径"
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题（可选）"
                    },
                    "content": {
                        "type": "string",
                        "description": "正文内容，多行用 \\n 分隔（可选）"
                    },
                    "filename": {
                        "type": "string",
                        "description": "文件名（不含扩展名，默认「文档」）（可选）"
                    }
                },
                "required": ["output_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": "读取 Word 文档（.docx）的内容，包括段落和表格",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": ".docx 文件路径"
                    }
                },
                "required": ["path"]
            }
        }
    }
]

HANDLERS = {
    "create_docx": create_docx,
    "read_docx": read_docx,
}


def register():
    for s in SCHEMAS:
        name = s["function"]["name"]
        register_tool(s, HANDLERS[name])
